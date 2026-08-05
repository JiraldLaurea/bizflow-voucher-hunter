"""
Builds the client-facing "How Voucher Hunt Works" Word document.

The document is a deliverable, not source: it is regenerated from here so the
wording stays reviewable in git and cannot drift silently from the product.

    pip install python-docx
    python scripts/build-how-it-works-doc.py

Screenshots are optional. Drop PNGs into docs/images/ using the filenames in
SHOTS below and they are embedded automatically; anything missing leaves a
short placeholder line instead.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
IMAGES = ROOT / "docs" / "images"
OUTPUT = ROOT / "docs" / "How Voucher Hunt Works.docx"

INK = RGBColor(0x0B, 0x1D, 0x3A)
MUTED = RGBColor(0x5B, 0x66, 0x7A)
PURPLE = RGBColor(0x5C, 0x3D, 0xFF)

SHOTS = {
    "directory": ("customer-01-directory.png", "The Home tab, showing live campaigns"),
    "campaign": ("customer-02-campaign.png", "A campaign page, ready to start"),
    "roulette": ("customer-03-roulette.png", "The reel mid-spin"),
    "datetime": ("customer-04-datetime.png", "Choosing a date and time slot"),
    "voucher": ("customer-05-voucher.png", "The issued voucher, with its QR code"),
    "validate": ("staff-01-validate.png", "Staff Validation, with a code entered"),
    "awarded": ("staff-02-awarded.png", "The 5% confirmed at the counter"),
}


def shot(doc, key):
    """Embeds the screenshot if it exists, else leaves a light placeholder."""
    filename, caption = SHOTS[key]
    path = IMAGES / filename
    if path.exists():
        doc.add_picture(str(path), width=Inches(2.6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        line = doc.add_paragraph(caption)
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = line.runs[0]
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED
        return
    line = doc.add_paragraph(f"[ Screenshot: {caption} — docs/images/{filename} ]")
    run = line.runs[0]
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def body(doc, text, bold_lead=None):
    paragraph = doc.add_paragraph()
    if bold_lead:
        lead = paragraph.add_run(bold_lead)
        lead.bold = True
    paragraph.add_run(text)
    return paragraph


def step(doc, number, title):
    heading = doc.add_heading(f"Step {number} — {title}", level=2)
    heading.runs[0].font.color.rgb = INK
    return heading


def build():
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(8)

    title = doc.add_heading("Voucher Hunt", level=0)
    title.runs[0].font.color.rgb = PURPLE
    subtitle = doc.add_paragraph(
        "How it works — a short guide for businesses considering the app"
    )
    subtitle.runs[0].font.size = Pt(13)
    subtitle.runs[0].font.color.rgb = MUTED

    body(
        doc,
        "This covers the two things that happen every day: a customer winning "
        "and booking a voucher, and your staff accepting it at the counter. No "
        "technical knowledge needed, and about five minutes to read.",
    )

    doc.add_heading("The idea in one paragraph", level=1)
    body(
        doc,
        "Customers open the app, pick your campaign, and spin for a discount — "
        "20% off, 50% off, a free dessert, whatever you decide. Winning is only "
        "half of it: to keep the voucher they must book a specific date and "
        "time, so the discount fills the tables you choose, not your Friday "
        "night rush. When they arrive and pay, your staff scan the code, and the "
        "customer automatically earns Loyalty Points worth 5% of what they "
        "spent — which they can spend back with you later.",
    )

    # ---- Part 1 -----------------------------------------------------------
    doc.add_page_break()
    doc.add_heading("Part 1 — The customer journey", level=1)

    step(doc, 1, "Find a campaign")
    body(
        doc,
        "The customer opens the app and sees every live campaign: your business "
        "name, your photo, your location, and the dates the offer runs. They can "
        "search or filter by category — Restaurant, Online Shop, Beauty, and so "
        "on. Campaigns that are fully booked stay visible but are marked as "
        "such, so your business keeps its presence even on a busy week.",
    )
    shot(doc, "directory")

    step(doc, 2, "Sign in with a mobile number")
    body(
        doc,
        "First time only. The customer enters their mobile number and types in "
        "the six-digit code we text them. No passwords, no account to remember — "
        "and it ties every voucher to a real, reachable person.",
    )

    step(doc, 3, "Start the hunt")
    body(
        doc,
        "The campaign page shows the offer, the rules, your address, a map, and "
        "a phone number to call you. One tap begins the hunt. Each customer gets "
        "three spins per campaign, and can earn a few more by sharing a referral "
        "link — capped daily, so it stays a game rather than a loophole.",
    )
    shot(doc, "campaign")

    step(doc, 4, "Spin and win")
    body(
        doc,
        "The reel spins and lands on a real prize from your campaign — say 30% "
        "off. The odds are yours to set: make the deep discounts rare and the "
        "everyday ones common.",
    )
    for text, lead in [
        (
            "If a tier has run out, or has no bookable time slots left, it is "
            "quietly taken out of the draw. Nobody wins something they cannot use.",
            "The app never offers a prize it cannot honour. ",
        ),
        (
            "They can keep the prize in hand, or spend another spin chasing "
            "better and risk ending up with less.",
            "The customer chooses when to stop. ",
        ),
    ]:
        bullet = doc.add_paragraph(style="List Bullet")
        run = bullet.add_run(lead)
        run.bold = True
        bullet.add_run(text)
    shot(doc, "roulette")

    step(doc, 5, "Pick a date and time")
    body(
        doc,
        "This is the part that makes the offer work for you. The customer picks "
        "from the slots you opened — the quiet Tuesday lunch, the 2pm off-peak "
        "table — and each slot has a set capacity. When it is full, it shows as "
        "sold out. You can also restrict your best prizes to your quietest "
        "slots: a 90% off tier might only ever be bookable at 2pm on a weekday.",
    )
    shot(doc, "datetime")

    step(doc, 6, "Confirm")
    body(doc, "The customer confirms and immediately receives:")
    for line in [
        "a voucher code (for example BIZ-6C1927) and a QR code in the app,",
        "a confirmation SMS with the discount, the date and the time,",
        "a reminder in the app's Vouchers tab, where it stays until it is used.",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    body(
        doc,
        "The seat is now reserved and the prize is deducted from your stock. One "
        "voucher per customer, per campaign.",
    )
    shot(doc, "voucher")

    # ---- Part 2 -----------------------------------------------------------
    doc.add_page_break()
    doc.add_heading("Part 2 — At the counter", level=1)
    body(
        doc,
        "Your staff need a phone, tablet or laptop, the dashboard, and a "
        "four-digit PIN. There is nothing to install.",
    )

    step(doc, 1, "Take the code")
    body(
        doc,
        "The customer shows their QR code, or reads out the voucher code. Your "
        "staff scan it or type it into Staff Validation. The screen immediately "
        "shows whether it is genuine: the discount, the customer's name, the "
        "booked date and time, and whether it has already been used. An expired "
        "or already-used voucher is refused on the spot.",
    )
    shot(doc, "validate")

    step(doc, 2, "Enter the amount they paid")
    body(
        doc,
        "Your staff apply the discount as normal on your own till, then enter "
        "the amount the customer actually paid and confirm. That is the whole "
        "job — one screen, two fields.",
    )

    step(doc, 3, "The 5% is awarded automatically")
    body(
        doc,
        "The moment the voucher is marked as used, the app works out 5% of what "
        "was paid and credits it to the customer's Loyalty Points wallet. "
        "₱1,000 paid becomes 50 LP. Your staff see the confirmation on screen, "
        "and the customer sees the new balance in their app.",
    )
    for line in [
        "No second scan, and nothing for your staff to remember.",
        "A customer earning for the first time gets a wallet automatically.",
        "The same sale can never be counted twice, however often the screen is retried.",
    ]:
        doc.add_paragraph(line, style="List Bullet")
    shot(doc, "awarded")

    # ---- Loyalty Points ---------------------------------------------------
    doc.add_page_break()
    doc.add_heading("What Loyalty Points mean for you", level=1)
    body(
        doc,
        "Loyalty Points are a shared currency across every partner in the "
        "network, and they bring customers back.",
    )
    for lead, text in [
        ("When you award points, ", "you owe the network their peso value — 50 LP issued costs you ₱50."),
        ("When a customer spends points with you, ", "the network owes you — 500 LP spent earns you ₱500."),
        (
            "At the end of each month the two are netted against each other. ",
            "If you owe more than you are owed, the difference comes out of your "
            "deposit. If you are owed more, we pay you the difference, less a 10% "
            "service fee.",
        ),
    ]:
        bullet = doc.add_paragraph(style="List Bullet")
        run = bullet.add_run(lead)
        run.bold = True
        bullet.add_run(text)
    body(
        doc,
        "Every partner keeps a ₱5,000 minimum deposit with the network. Your "
        "dashboard shows the running total at any time — what you owe, what you "
        "are owed, and the deposit behind both. You can also list items "
        "customers buy with points instead of money: a dessert for 150 LP, a "
        "facial for 1,200 LP. They pay in the app and collect at your counter, "
        "and the network pays you the peso value.",
    )

    doc.add_heading("What you get out of it", level=1)
    rows = [
        ("Fill the quiet hours", "Discounts are only bookable in the slots you open."),
        ("No wasted stock", "You set how many of each prize exists. When it is gone, it is gone."),
        ("Real bookings", "Every voucher is tied to a verified mobile number and a specific time."),
        ("They come back", "Points earned with you are spendable across the network — and with you."),
        ("You see everything", "Redemptions, no-shows, sales value, and what you owe or are owed."),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light List Accent 1"
    for name, detail in rows:
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = detail
        cells[0].paragraphs[0].runs[0].bold = True

    doc.add_heading("Common questions", level=1)
    faqs = [
        (
            "Can a customer use one voucher twice?",
            "No. It is marked used the moment your staff confirm it, and the app "
            "refuses it after that.",
        ),
        (
            "What if they do not show up?",
            "Staff can mark the booking a no-show. The slot is freed and it is "
            "recorded against that customer.",
        ),
        (
            "What if a customer wins something we cannot serve that day?",
            "Close the slot or the tier in your dashboard and it stops being "
            "offered immediately.",
        ),
        (
            "Do we need new hardware?",
            "No. Any phone, tablet or laptop with a browser.",
        ),
        (
            "What if we forget to enter the amount paid?",
            "The voucher is still marked used, but no Loyalty Points are awarded "
            "for that sale. The amount is what the 5% is calculated from.",
        ),
    ]
    for question, answer in faqs:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(question)
        run.bold = True
        doc.add_paragraph(answer)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    embedded = sum(1 for key in SHOTS if (IMAGES / SHOTS[key][0]).exists())
    print(f"Wrote {OUTPUT} ({embedded}/{len(SHOTS)} screenshots embedded)")


if __name__ == "__main__":
    build()
